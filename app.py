from flask import Flask, render_template, request, redirect, url_for, session, send_file, render_template_string, flash
from flask_sqlalchemy import SQLAlchemy
from models import db, Customer, Booking, Tables, MenuItem, OrderItem, User
from flask_admin import Admin
from flask_admin.contrib.sqla import ModelView
from io import BytesIO
from docx import Document
from openpyxl import Workbook
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfbase.ttfonts import TTFont
from datetime import timedelta
from dotenv import load_dotenv
import os

load_dotenv()

app = Flask(__name__)
app.secret_key = os.getenv('SECRET_KEY')
app.permanent_session_lifetime = timedelta(days=7)

app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///new_booking.db'
db.init_app(app)
admin = Admin(app, name='Misto Admin')


admin.add_view(ModelView(User, db.session))
admin.add_view(ModelView(MenuItem, db.session))

@app.before_request
def load_logged_in_user():
    user_id = session.get('user_id')
    if user_id is None:
        g.user = None
        g.customer = None
        g.bookings = []
    else:
        g.user = User.query.get(user_id)
        g.customer = Customer.query.filter_by(email=g.user.email).first()
        g.bookings = Booking.query.filter_by(customer_id=g.customer.id).all() if g.customer else []


@app.route('/export_booking/<format>')
def export_booking(format):
    bookings = Booking.query.all()

    if format == 'docx':
        doc = Document()
        doc.add_heading('Список бронювань', 0)

        for b in bookings:
            customer = Customer.query.get(b.customer_id)
            order_items = OrderItem.query.filter_by(booking_id=b.id).all()
            dishes = []
            for item in order_items:
                menu_item = MenuItem.query.get(item.menu_item_id)
                if menu_item:
                    dishes.append(menu_item.name)
            dishes_str = ", ".join(dishes) if dishes else "немає"

            doc.add_paragraph(
                f"Ім’я: {customer.name}, Телефон: {customer.phone}, Email: {customer.email}, "
                f"Дата: {b.day}, Час: {b.time}, Гостей: {b.guests}, Стіл: {b.table_id}, Страви: {dishes_str}"
            )

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name='booking.docx')

    elif format == 'xlsx':
        wb = Workbook()
        ws = wb.active
        ws.append(['Ім’я', 'Телефон', 'Email', 'Дата', 'Час', 'Гостей', 'Тривалість', 'Стіл', 'Страви'])

        for b in bookings:
            customer = Customer.query.get(b.customer_id)
            order_items = OrderItem.query.filter_by(booking_id=b.id).all()
            dishes = []
            for item in order_items:
                menu_item = MenuItem.query.get(item.menu_item_id)
                if menu_item:
                    dishes.append(menu_item.name)
            dishes_str = ", ".join(dishes) if dishes else "немає"

            ws.append([
                customer.name, customer.phone, customer.email,
                b.day, b.time, b.guests, b.during, b.table_id, dishes_str
            ])

        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name='booking.xlsx')

    elif format == 'pdf':
        buffer = BytesIO()

        doc = SimpleDocTemplate(buffer, pagesize=A4,
                                leftMargin=2 * cm, rightMargin=2 * cm,
                                topMargin=2 * cm, bottomMargin=2 * cm)

        font_path = os.path.join('static', 'fonts', 'DejaVuSans.ttf')
        pdfmetrics.registerFont(TTFont('DejaVuSans', font_path))

        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='MyStyle', fontName='DejaVuSans', fontSize=10, leading=14))
        styles.add(ParagraphStyle(name='MyHeading', fontName='DejaVuSans', fontSize=16, leading=20, alignment=1))

        content = []
        content.append(Paragraph("Список бронювань", styles['MyHeading']))
        content.append(Spacer(1, 12))

        for b in bookings:
            customer = Customer.query.get(b.customer_id)
            order_items = OrderItem.query.filter_by(booking_id=b.id).all()
            dishes = []
            for item in order_items:
                menu_item = MenuItem.query.get(item.menu_item_id)
                if menu_item:
                    dishes.append(menu_item.name)
            dishes_str = ", ".join(dishes) if dishes else "немає"

            text = f"""
            <b>Ім’я:</b> {customer.name}, <b>Телефон:</b> {customer.phone}, <b>Email:</b> {customer.email}<br/>
            <b>Дата:</b> {b.day}, <b>Час:</b> {b.time}, <b>Гостей:</b> {b.guests}, <b>Тривалість:</b> {b.during}, <b>Стіл:</b> {b.table_id}<br/>
            <b>Страви:</b> {dishes_str}
            """
            content.append(Paragraph(text, styles['MyStyle']))
            content.append(Spacer(1, 8))

        doc.build(content)
        buffer.seek(0)

        return send_file(buffer, as_attachment=True, download_name="booking.pdf")


@app.route('/')
def index():
    user = None
    if 'user_id' in session:
        user = User.query.get(session['user_id'])
    return render_template('main.html', user=user)

@app.route('/profile')
def profile():
    if not g.user:
        return redirect(url_for('authorize'))
   
    return render_template('profile.html')


@app.route('/menu')
def menu():
    user = None
    if 'user_id' in session:
        user = User.query.get(session['user_id'])
    menu_items = [
        {
            "id": 1,
            "title": "Стейк зі свинини на грилі зі смаженою картоплею",
            "price": 289,
            "image_url": "/static/uploads/3l0GvTy9P8lMAjho8bdBK4f37C9qqdyJg18zX5fX.webp"
        },
        {
            "id": 2,
            "title": "Млинці з м'ясом та томатним соусом",
            "price": 188,
            "image_url": "/static/uploads/7fYwur0eTaKgc51iLkptgxemwBML8qk82KlExi7V.webp"
        },
        {
            "id": 3,
            "title": "Паста Карбонара",
            "price": 176,
            "image_url": "static/uploads/Lhm6HJbNA8eDlzMHQuJT2Yit8gcEo7FLxi1YW0o6.webp"
        },
        {
            "id": 4,
            "title": "Стейк лосося з кус-кусом в соусі Песто",
            "price": 289,
            "image_url": "static/uploads/URRmjVgmx3lhwvcZFNRNDHk2JCxGlcyiSVgI8JLz.webp"
        },
        {
            "id": 5,
            "title": "Паста у томатному соусі",
            "price": 210,
            "image_url": "static/uploads/xKwzERFOQOAfILBh0lvpLaW1mtl6fxEZWXcfQUq2.webp"
        },
        {
            "id": 6,
            "title": "Салат Грецький",
            "price": 144,
            "image_url": "static/uploads/NAwZlcM-TYIbKyk-gsKwehF.webp"
        },
        {
            "id": 7,
            "title": "Оладок з цукіні, яйцем пашот та слабосоленою фореллю",
            "price": 490,
            "image_url": "static/uploads/pancake_salmon.jpg"
        },
        {
            "id": 8,
            "title": "Скрембл з мортаделою",
            "price": 430,
            "image_url": "static/uploads/scramble_avocado.webp"
        },
        {
            "id": 9,
            "title": "Скрембл з лососем",
            "price": 445,
            "image_url": "static/uploads/Scrambled-Eggs-with-Anchovy-Butter-Secondary-Image.webp"
        },
        {
            "id": 10,
            "title": "Англійський сніданок",
            "price": 320,
            "image_url": "static/uploads/english.webp"
        },
        {
            "id": 11,
            "title": "Сирники з солоною карамеллю та сметаною",
            "price": 280,
            "image_url": "static/uploads/cheesecake.webp"
        },

        {
            "id": 12,
            "title": "Coca-cola",
            "price": 54,
            "image_url": "static/uploads/FvHplte-bbCyvnk-ZJjreGh.webp"
        },
        {
            "id": 13,
            "title": "Вода Bonaqua",
            "price": 44,
            "image_url": "static/uploads/pCuMoDz-MvqddzF-fzodjjw.webp"
        },
        {
            "id": 14,
            "title": "Sprite",
            "price": 56,
            "image_url": "static/uploads/sprite.webp"
        }, 
        {
            "id": 15,
            "title": "Чай",
            "price": 44,
            "image_url": "static/uploads/tea.webp"
        }, 
        {
            "id": 16,
            "title": "Еспрессо",
            "price": 50,
            "image_url": "static/uploads/espresso.jpg"
        }
        
    ]
    return render_template('menu.html', menu_items=menu_items, user=user)

@app.route('/add_to_cart', methods=['POST'])
def add_to_cart():
    item_id = int(request.form.get('item_id'))
    title = request.form.get('title')
    price = float(request.form.get('price'))

    cart = session.get('cart', [])

    for item in cart:
        if item['item_id'] == item_id:
            item['quantity'] += 1
            break
    else:
        cart.append({'item_id': item_id, 'title': title, 'price': price, 'quantity': 1})

    session['cart'] = cart
    session.modified = True

    return redirect(url_for('menu'))

@app.route('/cart')
def cart():
    user = None
    if 'user_id' in session:
        user = User.query.get(session['user_id'])
    cart_items = session.get('cart', [])
    return render_template('cart.html', cart=cart_items, user=user)

@app.route('/remove_from_cart/<int:item_index>', methods=['POST'])
def remove_from_cart(item_index):
    cart = session.get('cart', [])
    if 0 <= item_index < len(cart):
        cart.pop(item_index)
    session['cart'] = cart
    session.modified = True
    return redirect(url_for('cart'))


@app.route('/authorize', methods=['GET', 'POST'])
def authorize():
    if request.method == 'POST':
        username = request.form['username'].strip()
        password = request.form['password'].strip()

        user = User.query.filter_by(username=username, password=password).first()
        if not user:
            flash("Невірний логін або пароль")
            return redirect(url_for('authorize'))

        session.clear()
        session['user_id'] = user.id
        session['role'] = user.role.lower().strip()
        session.permanent = True

        if session['role'] == 'admin':
            return redirect(url_for('show'))
        else:
            return redirect(url_for('profile'))

    return render_template('authorize.html')


@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username'].strip()
        email = request.form['email'].strip()
        password = request.form['password']

        if not username or not password:
            print("Будь ласка, введіть ім'я користувача та пароль")
            return redirect(url_for('register'))

        existing_user = User.query.filter_by(username=username).first()
        if existing_user:
            print("Користувач вже існує!")
            return redirect(url_for('register'))

        role = request.form.get('role', 'user') 
        new_user = User(username=username, email=email, password=password, role=role)
        db.session.add(new_user)
        db.session.commit()

        return redirect(url_for('authorize'))

    return render_template('register.html')

@app.route('/home')
def home():
    if 'user' not in session:
        return redirect(url_for('authorize'))
    return render_template('main.html', user=session['user'])

@app.route('/booking', methods=['GET', 'POST'])
def booking():
    user = None
    if 'user_id' in session:
        user = User.query.get(session['user_id'])
    if request.method == 'POST':
        session['day'] = request.form.get('day')
        session['time'] = request.form.get('time')
        session['guests'] = request.form.get('guests')
        session['during'] = request.form.get('during')
        session['table_id'] = request.form.get('table_id')
        return redirect(url_for('contacts'))

    cart_items = session.get('cart', [])
    tables = Tables.query.all()
    
    return render_template('booking.html', tables=tables, cart=cart_items, user=user)


@app.route('/contacts')
def contacts():
    user = None
    if 'user_id' in session:
        user = User.query.get(session['user_id'])
    return render_template('contacts.html', user=user)

@app.route('/add_contacts', methods=['POST'])
def add_contacts():
    for field in ['day', 'time', 'guests', 'during', 'table_id']:
        if field not in session:
            return redirect(url_for('booking'))

    if 'user_id' in session:
        user = User.query.get(session['user_id'])
        email = user.email
        name = request.form.get('name') or user.username
        phone = request.form.get('tel')
    else:
        email = request.form['email']
        name = request.form['name']
        phone = request.form['tel']

    
    customer = Customer.query.filter_by(email=email).first()
    if not customer:
        customer = Customer(name=name, email=email, phone=phone)
        db.session.add(customer)
        db.session.commit()

    booking = Booking(
        day=session['day'],
        time=session['time'],
        guests=session['guests'],
        during=session['during'],
        table_id=session['table_id'],
        customer_id=customer.id
    )
    db.session.add(booking)
    db.session.commit()

    cart = session.get('cart', [])
    for item in cart:
        menu_item = MenuItem.query.get(item['item_id'])
        if menu_item:
            order_item = OrderItem(
            booking_id=booking.id,
            menu_item_id=menu_item.id,
            quantity=item['quantity']
        )
        db.session.add(order_item)
    db.session.commit()

    return redirect(url_for('confirmation', booking_id=booking.id))

@app.route('/confirmation/<int:booking_id>')
def confirmation(booking_id):
    booking = Booking.query.get_or_404(booking_id)
    customer = Customer.query.get(booking.customer_id)
    table = Tables.query.get(booking.table_id)

    if not customer:
        return redirect(url_for('profile'))
    
    user = None
    if 'user_id' in session:
        user = User.query.get(session['user_id'])

    return render_template(
        'booking_result.html',
        booking=booking,
        name=customer.name,
        email=customer.email,
        table=table,
        user=user
    )
    

@app.route('/show')
def show():
    if 'user_id' not in session or session.get('role') != 'admin':
        return redirect(url_for('authorize'))

    query = Booking.query

    name = request.args.get('name')
    day = request.args.get('day')
    table_id = request.args.get('table_id')

    if name:
        query = query.join(Customer).filter(Customer.name.ilike(f'%{name}%'))
    if day:
        query = query.filter(Booking.day == day)
    if table_id:
        query = query.filter(Booking.table_id == table_id)

    customers = {c.id: c for c in Customer.query.all()}
    order_items = OrderItem.query.all()
    menu_items = {m.id: m for m in MenuItem.query.all()}

    orders_by_booking = {}
    for item in order_items:
        menu_item = menu_items.get(item.menu_item_id)
        if menu_item:
            orders_by_booking.setdefault(str(item.booking_id), []).append({
                'title': menu_item.name,
                'quantity': item.quantity,
                'price':  menu_item.price
            })

    bookings = query.all()
    tables = Tables.query.all()

    return render_template('show.html', reservations=bookings, tables=tables, orders=orders_by_booking,  customers=customers,)


@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('index'))

@app.before_request
def init_menu_items():
    if not MenuItem.query.first(): 
        db.session.add_all([
            MenuItem(id=1, name="Стейк зі свинини на грилі зі смаженою картоплею", price=289),
            MenuItem(id=2, name="Млинці з м'ясом та томатним соусом", price=188),
            MenuItem(id=3, name="Паста Карбонара", price=176),
            MenuItem(id=4, name="Стейк лосося з кус-кусом в соусі Песто", price=289),
            MenuItem(id=5, name="Паста у томатному соусі", price=210),
            MenuItem(id=6, name="Салат Грецький", price=144),
            MenuItem(id=7, name="Оладок з цукіні, яйцем пашот та слабосоленою фореллю", price=490),
            MenuItem(id=8, name="Скрембл з мортаделою", price=430),
            MenuItem(id=9, name="Скрембл з лососем", price=445),
            MenuItem(id=10, name="Англійський сніданок", price=320),
            MenuItem(id=11, name="Сирники з солоною карамеллю та сметаною", price=280),
            MenuItem(id=12, name="Coca-cola", price=54),
            MenuItem(id=13, name="Вода Bonaqua", price=44),
            MenuItem(id=14, name="Sprite", price=56),
            MenuItem(id=15, name="Чай", price=44),
            MenuItem(id=16, name="Еспрессо", price=50),
        ])
        db.session.commit()


@app.route('/delete_booking/<int:booking_id>', methods=['POST'])
def delete_booking(booking_id):
    if 'user_id' not in session or session.get('role') != 'admin':
        return redirect(url_for('authorize'))

    booking = Booking.query.get_or_404(booking_id)
    db.session.delete(booking)
    db.session.commit()
    return redirect(url_for('show'))


@app.before_request
def init_tables():
    if not Tables.query.first():
        db.session.add_all([
            Tables(table_number="1", seats=2),
            Tables(table_number="2", seats=4),
            Tables(table_number="3", seats=6),
            Tables(table_number="4", seats=8),
        ])
        db.session.commit()
        print("Столи успішно додані!")

if __name__ == "__main__":
    with app.app_context():
        db.create_all()
        init_tables()
    app.run(debug=True)
